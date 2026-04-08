"""Shared fixtures for gateway integration tests."""
import io

from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

GATEWAY_URL = "http://gateway:8000"
TIMEOUT = 120


def make_png(text: str = "Hello World", width: int = 200, height: int = 100) -> bytes:
    """Generate a PNG image with text drawn on it."""
    img = Image.new("RGB", (width, height), color="white")
    draw = ImageDraw.Draw(img)
    draw.text((10, 40), text, fill="black")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def make_pdf(text: str = "This is a test PDF document.", pages: int = 1) -> bytes:
    """Generate a simple PDF with text."""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    for i in range(pages):
        c.drawString(72, 700, f"{text} Page {i + 1}.")
        c.showPage()
    c.save()
    return buf.getvalue()


def make_pdf_with_image(text: str = "Document with figure below.") -> bytes:
    """Generate a PDF containing text and an embedded image."""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.drawString(72, 750, text)

    img_bytes = make_png("Chart Data")
    img_buf = io.BytesIO(img_bytes)
    c.drawImage(img_buf, 72, 500, width=200, height=100)

    c.showPage()
    c.save()
    return buf.getvalue()


def make_docx(text: str = "This is a test Word document.") -> bytes:
    """Generate a simple DOCX."""
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph(text)
    doc.add_heading("Section Two", level=2)
    doc.add_paragraph("More content here.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_xlsx() -> bytes:
    """Generate a simple XLSX with a table."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Test"
    ws.append(["Name", "Value"])
    ws.append(["Alpha", 10])
    ws.append(["Beta", 20])
    ws.append(["Gamma", 30])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
