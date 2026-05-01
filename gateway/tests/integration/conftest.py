"""Shared fixtures for gateway integration tests."""
import io
import os
from collections import namedtuple

from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas


def pytest_collection_modifyitems(config, items):
    """When FOCUS_TESTS=1, run only tests marked @pytest.mark.focus."""
    if os.environ.get("FOCUS_TESTS") != "1":
        return
    focus_items = [item for item in items if item.get_closest_marker("focus")]
    if focus_items:
        items[:] = focus_items

EmbeddedImage = namedtuple("EmbeddedImage", ["label", "width", "height"])

GATEWAY_URL = "http://gateway:8000"
TIMEOUT = 120


def make_png(text: str = "Hello World", width: int = 200, height: int = 100) -> bytes:
    """Generate a PNG image with text drawn on it."""
    img = Image.new("RGB", (width, height), color="white")
    draw = ImageDraw.Draw(img)
    draw.text((10, 40), text, fill="black")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def make_pdf(text: str = "This is a test PDF document.", pages: int = 1) -> bytes:
    """Generate a simple PDF with text."""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    for i in range(pages):
        c.drawString(72, 700, f"{text} Page {i + 1}.")
        c.showPage()
    c.save()
    return buf.getvalue()


def make_pdf_with_image(text: str = "Document with figure below.") -> bytes:
    """Generate a PDF containing text and an embedded image."""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.drawString(72, 750, text)

    img_bytes = make_png("Chart Data")
    c.drawImage(ImageReader(io.BytesIO(img_bytes)), 72, 500, width=200, height=100)

    c.showPage()
    c.save()
    return buf.getvalue()


def make_pdf_with_images(images: list[EmbeddedImage], text: str = "Document with embedded images.") -> bytes:
    """Generate a PDF containing text and multiple embedded images of varying sizes."""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.drawString(72, 800, text)

    y_cursor = 750
    for img in images:
        img_bytes = make_png(img.label, width=img.width, height=img.height)
        c.drawImage(ImageReader(io.BytesIO(img_bytes)), 72, y_cursor - img.height, width=img.width, height=img.height)
        y_cursor -= img.height + 20
        if y_cursor < 100:
            c.showPage()
            y_cursor = 750

    c.showPage()
    c.save()
    return buf.getvalue()


def make_tiff(pages: list[str]) -> bytes:
    """Generate a multi-page TIFF with text drawn on each page."""
    frames = []
    for text in pages:
        img = Image.new("RGB", (200, 100), color="white")
        draw = ImageDraw.Draw(img)
        draw.text((10, 40), text, fill="black")
        frames.append(img)
    buf = io.BytesIO()
    frames[0].save(buf, format="TIFF", save_all=True, append_images=frames[1:])
    return buf.getvalue()


def make_docx(text: str = "This is a test Word document.") -> bytes:
    """Generate a simple DOCX."""
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph(text)
    doc.add_heading("Section Two", level=2)
    doc.add_paragraph("More content here.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_xlsx() -> bytes:
    """Generate a simple XLSX with a table."""
    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "Test"
    ws.append(["Name", "Value"])
    ws.append(["Alpha", 10])
    ws.append(["Beta", 20])
    ws.append(["Gamma", 30])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
